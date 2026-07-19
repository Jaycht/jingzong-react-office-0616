import os, re, json
from docx import Document

SRC = [
    "D:/下载/关于办理商业贿赂刑事案件适用法律 若干问题的意见.docx",
    "D:/下载/最高人民法院、最高人民检察院关于办理走私刑事案件适用法律若干问题的解释_20140812.docx",
]
OUT = ".import_tmp"

ART_RE = re.compile(r'^第([一二三四五六七八九十百零〇两]+)条(之一)?[　 \t]*(.*)$')

def extract(path):
    name = os.path.splitext(os.path.basename(path))[0]
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for ti, tbl in enumerate(doc.tables):
        parts.append(f"\n[表格 {ti+1}]")
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts)
    # 清理多余空行
    lines = [ln.rstrip() for ln in text.split("\n")]
    text = "\n".join(lines)
    outp = os.path.join(OUT, name + ".txt")
    with open(outp, "w", encoding="utf-8") as f:
        f.write(text)
    arts = [l for l in lines if ART_RE.match(l)]
    print(f"=== {name} ===")
    print(f"  字符数: {len(text)}")
    print(f"  条数(第X条): {len(arts)}")
    if arts:
        print(f"  首条: {arts[0][:50]}")
        print(f"  末条: {arts[-1][:50]}")
    # 版本线索（前 12 行）
    head = [l for l in lines[:12] if l.strip()][:8]
    print(f"  头部: {head}")
    return name

for s in SRC:
    extract(s)
