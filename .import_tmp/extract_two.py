import os, re
from docx import Document

SRC = [
    r"D:/下载/公安机关办理刑事案件程序规定-2020年修正版.docx",
    r"D:/下载/公安机关办理行政案件程序规定.docx",
]
OUT = r".import_tmp"
ART_RE = re.compile(r'^第([一二三四五六七八九十百零〇两]+)条(之一)?[　 \t]*(.*)$')

def extract(path):
    d = Document(path)
    parts = []
    for p in d.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            line = ' '.join(c for c in cells if c)
            if line:
                parts.append(line)
    return '\n'.join(parts)

for src in SRC:
    base = os.path.splitext(os.path.basename(src))[0]
    clean = base.replace('-2020年修正版', '').strip()
    txt = extract(src)
    out_path = os.path.join(OUT, clean + '.txt')
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(txt)
    arts = [l for l in txt.splitlines() if ART_RE.match(l)]
    print('===', base)
    print('干净名:', clean)
    print('字符数:', len(txt))
    print('条数(行首中文条号):', len(arts))
    if arts:
        print('首条:', arts[0][:70])
        print('末条:', arts[-1][:70])
    print('--- 尾部 500 字(查版本/日期) ---')
    print(txt[-500:])
    print()
