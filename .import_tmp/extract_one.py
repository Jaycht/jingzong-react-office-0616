import os, sys, re, json
from docx import Document

SRC = "D:/下载/公安部刑事案件管辖分工规定.docx"
OUT = ".import_tmp/公安部刑事案件管辖分工规定.txt"

doc = Document(SRC)
parts = []
for p in doc.paragraphs:
    parts.append(p.text)
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        parts.append(" | ".join(cells))

text = "\n".join(parts)
text = re.sub(r"[ \t]+", " ", text)
text = re.sub(r"\n{3,}", "\n\n", text).strip()

os.makedirs(os.path.dirname(OUT), exist_ok=True)
with open(OUT, "w", encoding="utf-8") as f:
    f.write(text)

# 条文统计（中文数字条号）
art_re = re.compile(r"^第([一二三四五六七八九十百零〇两]+)条")
lines = text.split("\n")
art_lines = [l for l in lines if art_re.match(l.strip())]
# 取末条号
last = None
for l in art_lines:
    m = art_re.match(l.strip())
    if m: last = m.group(1)

def cn2int(s):
    d = {'零':0,'〇':0,'一':1,'二':2,'两':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9}
    if s.isdigit(): return int(s)
    total, sec, unit = 0, 0, 1
    for ch in s:
        if ch in d:
            sec += d[ch]
        elif ch == '十':
            unit = 10; sec = sec if sec else 1
        elif ch == '百':
            unit = 100; sec = sec if sec else 1
        elif ch == '千':
            unit = 1000; sec = sec if sec else 1
        total += sec * unit if ch in '十百千' else 0
        if ch in '十百千': sec = 0
    return total + sec

report = {
    "chars": len(text),
    "lines": len(lines),
    "article_lines": len(art_lines),
    "last_article_cn": last,
    "last_article_int": cn2int(last) if last else None,
    "head": lines[:8],
}
print(json.dumps(report, ensure_ascii=False, indent=2))
# 打印版本/施行线索
for kw in ["公安部令", "部务会议", "起施行", "印发", "发布", "修正"]:
    for l in lines[:40]:
        if kw in l:
            print("VER:", l.strip()[:120])
            break
