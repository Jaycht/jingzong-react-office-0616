# -*- coding: utf-8 -*-
import os, re
from docx import Document

SRC = "D:/下载/最高人民法院、最高人民检察院关于办理走私刑事案件适用法律若干问题的解释_20140812.docx"
OUT = ".import_tmp/关于办理走私刑事案件适用法律若干问题的解释.txt"

doc = Document(SRC)
lines = []
for p in doc.paragraphs:
    t = p.text.strip()
    if t:
        lines.append(t)
for tbl in doc.tables:
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        line = " | ".join(cells)
        if line.strip():
            lines.append(line)

body = "\n".join(lines)
arts = re.findall(r'^第([一二三四五六七八九十百零〇两]+)条', body, re.M)
print("字符数:", len(body))
print("条文数(第X条):", len(arts))
print("首段:", lines[:3])
print("末段:", lines[-3:])
open(OUT, "w", encoding="utf-8").write(body)
print("已写入:", OUT)
