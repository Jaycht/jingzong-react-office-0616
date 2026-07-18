import os, re, json, glob
from docx import Document

SRC = r"D:\下载"
OUT = r"E:\Deployment\WorkBuddy\jingzong-react-office-0616\.import_tmp"
os.makedirs(OUT, exist_ok=True)

ART_RE = re.compile(r"^第[一二三四五六七八九十百零〇两0-9]+条")

def extract_text(path):
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        t = p.text
        if t is not None:
            t = t.strip()
            if t:
                parts.append(t)
    # also pull table cell text (some laws embed tables)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    parts.append(ct)
    return "\n".join(parts)

report = []
for path in sorted(glob.glob(os.path.join(SRC, "*.docx"))):
    name = os.path.basename(path)
    if name.startswith("~$"):
        continue
    try:
        text = extract_text(path)
    except Exception as e:
        report.append({"file": name, "error": str(e)})
        continue
    n = len(text)
    # count article lines (starts with 第..条) and all 第X条 occurrences
    lines = text.split("\n")
    art_lines = [l for l in lines if ART_RE.match(l)]
    all_arts = len(re.findall(r"第[一二三四五六七八九十百零〇两0-9]+条", text))
    # safe out name = original without .docx
    base = name[:-5]
    outp = os.path.join(OUT, base + ".txt")
    with open(outp, "w", encoding="utf-8") as f:
        f.write(text)
    report.append({
        "file": name,
        "out": base + ".txt",
        "chars": n,
        "article_lines": len(art_lines),
        "all_art_occurrences": all_arts,
        "head": text[:120].replace("\n", " "),
    })

with open(os.path.join(OUT, "_report.json"), "w", encoding="utf-8") as f:
    json.dump(report, f, ensure_ascii=False, indent=2)

print("extracted:", len([r for r in report if "error" not in r]), "errors:", len([r for r in report if "error" in r]))
for r in report:
    if "error" in r:
        print("ERR", r["file"], r["error"])
    else:
        print(f'{r["chars"]:>7}字 | 条行:{r["article_lines"]:>3} 条次:{r["all_art_occurrences"]:>3} | {r["file"]}')
